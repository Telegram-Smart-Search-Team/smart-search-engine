import os
import sys
import shutil
import pathlib
import asyncio
import aiohttp
import tempfile
import typing as tp
from dotenv import load_dotenv

BASE_DIR = pathlib.Path(__file__).parents[2]
load_dotenv(dotenv_path=str(BASE_DIR / ".env"))

sys.path.append(str(BASE_DIR / "src"))

# pdf imports
import fitz  # PyMuPDF

# word imports
from docx import Document

# excel imports
from openpyxl import load_workbook

# local imports
from utils.utils import *


MASK_TEXT_DEFAULT_PROMPT = """
You are a careful sanitizer.

Return the SAME text, but with sensitive data masked. Preserve every character of the original text except the sensitive spans you mask (keep punctuation, quotes, whitespace, line breaks, ordering, and wording unchanged).

Mask with asterisks: replace each sensitive span with exactly "********" (8 asterisks). Do not add explanations. Output ONLY the sanitized text.

What counts as sensitive (mask these):
1) Credentials / secrets
- Passwords / passphrases / PINs (including “password=…”, “pwd: …”, “pass: …”, “pin: …”)
- API keys, access keys, secret keys, client secrets
- Tokens of any kind: Bearer tokens, JWTs, OAuth access/refresh tokens, session tokens, CSRF tokens
- Cookies / session ids (e.g., “Cookie: …”, “sessionid=…”, “auth=…”, “sid=…”)
- Private keys / key blocks / certificates / signatures (PEM blocks, SSH private keys, etc.)
- Database / broker / service connection strings that contain credentials (e.g., postgres://user:pass@…, amqp://…)
- Secrets in config/code/logs: .env, YAML/JSON, headers, query params, CLI flags

2) Personal/financial identifiers (PII/PCI)
- Emails, phone numbers, usernames/handles if clearly identifying in context
- Home addresses, exact location strings, full names if used as identifying data in a record
- Government/identity numbers (passport, national id, SSN-like, tax id)
- Payment data: card numbers, CVV/CVC, IBAN, bank account numbers

Detection rules (use multiple signals):
- Any value following keys/labels like: password, passwd, pwd, secret, token, api_key, apikey, key, auth, authorization, bearer, cookie, session, sid, jwt, private_key, client_secret, refresh_token, access_token
- Any high-entropy long string (typically 16+ chars) that looks like a secret (mixed case/digits, base64/hex-like, JWT “xxx.yyy.zzz”, etc.)
- URLs with credentials or secret query params (e.g., ?token=…, ?key=…, ?password=…)

Mask only the VALUE part (keep the key/label and surrounding syntax):
- "password=abc" -> "password=********"
- "Authorization: Bearer eyJ..." -> "Authorization: Bearer ********"
- "email: a@b.com" -> "email: ********"

If unsure, prefer masking over leaking.
""".strip()


async def mask_text(
    text: str,
    system_prompt: str = MASK_TEXT_DEFAULT_PROMPT,
):
    async with aiohttp.ClientSession() as session:
        payload = {
            "system": system_prompt,
            "prompt": text,
            "max_new_tokens": 4096,
            "temperature": 0.0,
        }

        async with session.post(os.getenv("INTERNAL_TEXT_ENDPOINT"), json=payload) as response:
            data = await response.json()
            return data["content"]  # str


async def describe_image(
    path: str,
    system_prompt: str = "You are an expert image describer for Telegram media attachments.",
    prompt: str = "Describe what happens on the image. Describe people, objects, location, actions, emotions, on-screen text and logos. Try to identify and name some well known people, brands or animated characters.",
):
    async with aiohttp.ClientSession() as session:
        with open(path, "rb") as file:
            form = aiohttp.FormData()
            form.add_field(
                "image_file", file, filename=os.path.basename(path), content_type=guess_mimetype(path, "image/png")
            )
            form.add_field("system", system_prompt)
            form.add_field("prompt", prompt)

            async with session.post(os.getenv("INTERNAL_IMAGE_ENDPOINT"), data=form) as response:
                data = await response.json()
                return data["content"]  # str


async def describe_video(
    path: str,
    system_prompt: str = "You are an expert video describer for Telegram media attachments.",
    prompt: str = "Describe what happens in this video. Split the clip into sensible segments. Describe people, objects, locations, actions, emotions, camera movement, on-screen text and logos. Try to identify and name some well known people, brands or animated characters.",
):
    async with aiohttp.ClientSession() as session:
        with open(path, "rb") as file:
            form = aiohttp.FormData()
            form.add_field(
                "video_file", file, filename=os.path.basename(path), content_type=guess_mimetype(path, "video/mp4")
            )
            form.add_field("system", system_prompt)
            form.add_field("prompt", prompt)

            async with session.post(os.getenv("INTERNAL_VIDEO_ENDPOINT"), data=form) as response:
                data = await response.json()
                return data["content"]  # str


async def transcribe_audio(path: str):
    async with aiohttp.ClientSession() as session:
        with open(path, "rb") as file:
            form = aiohttp.FormData()
            form.add_field(
                "audio_file", file, filename=os.path.basename(path), content_type=guess_mimetype(path, "audio/mpeg")
            )

            async with session.post(os.getenv("INTERNAL_AUDIO_ENDPOINT"), data=form) as response:
                data = await response.json()
                return data[
                    "segments"
                ]  # [(seg["start"]: float, seg["end"]: float, seg["content"]: str) for seg in segments]


async def transcribe_pdf(path: str, image_max_pages: int = 3) -> str:
    # normal text extraction
    text_parts: list[str] = []

    try:
        with fitz.open(path) as doc:
            for page in doc:
                txt = page.get_text("text")
                if txt:
                    text_parts.append(txt)
    except Exception:
        # pdf seems to be corrupted or unreadable
        text_parts = []

    plain_text = "\n".join(text_parts).strip()
    if plain_text:
        return plain_text

    # describing as images
    if image_max_pages <= 0:
        return ""

    tmp_dir = pathlib.Path(tempfile.mkdtemp(prefix="pdf_pages_"))

    try:
        with fitz.open(path) as doc:
            num_pages = min(image_max_pages, len(doc))

            tasks: list[asyncio.Task] = []

            for i in range(num_pages):
                page = doc.load_page(i)
                pix = page.get_pixmap(dpi=200)

                img_path = tmp_dir / f"{pathlib.Path(path).stem}_page_{i + 1}.png"
                pix.save(str(img_path))

                tasks.append(
                    asyncio.create_task(
                        describe_image(
                            str(img_path),
                            system_prompt="You are an expert PDF transcriber.",
                            prompt="Transcribe the text you see on this PDF page. Detect the language for each section, and transcribe every part in its original language. Output only the text itself.",
                        )
                    )
                )

            # calls concurrently
            page_descriptions = await asyncio.gather(*tasks, return_exceptions=True)

        result_chunks: list[str] = []
        for index, description in enumerate(page_descriptions, start=1):
            if isinstance(description, Exception):
                continue

            if description and description.strip():
                result_chunks.append(f"=== Page {index} ===\n{description.strip()}")

        return "\n\n".join(result_chunks).strip()

    finally:
        # clean up
        shutil.rmtree(tmp_dir, ignore_errors=True)


def transcribe_docx(path: str) -> str:
    doc = Document(path)
    parts: list[str] = []

    # paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    # tables (cell text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                parts.append("\t".join(row_text))

    return "\n".join(parts)


def transcribe_xlsx(path: str) -> str:
    wb = load_workbook(filename=path, data_only=True)
    parts: list[str] = []

    for sheet in wb.worksheets:
        parts.append(f"=== SHEET: {sheet.title} ===")
        for row in sheet.iter_rows():
            row_vals = []
            for cell in row:
                if cell.value is None:
                    row_vals.append("")
                else:
                    row_vals.append(str(cell.value))

            # skipping completely empty rows
            if any(val.strip() for val in row_vals):
                parts.append("\t".join(row_vals))

    return "\n".join(parts)


async def file_to_text(path: str) -> str:
    p = pathlib.Path(path)

    if not p.exists():
        raise FileNotFoundError(f"File does not exist: {path}")

    ext = p.suffix.lower()

    pdf_exts = {".pdf"}
    word_exts = {".docx"}
    excel_exts = {".xlsx", ".xlsm", ".xltx", ".xltm"}
    image_exts = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"}
    video_exts = {".mp4", ".mov", ".mkv", ".avi", ".webm", ".flv", ".m4v", ".gif"}
    audio_exts = {".mp3", ".wav", ".ogg", ".m4a", ".flac", ".aac", ".wma", ".opus"}
    text_exts = {".txt", ".md", ".log"}

    # pdf documents
    if ext in pdf_exts:
        return await transcribe_pdf(str(p))

    # word documents
    if ext in word_exts:
        return transcribe_docx(str(p))

    # excel documents
    if ext in excel_exts:
        return transcribe_xlsx(str(p))

    # image files
    if ext in image_exts:
        return await describe_image(str(p))

    # video files
    if ext in video_exts:
        return await describe_video(str(p))

    # audio files
    if ext in audio_exts:
        # segments are expected to be a list like:
        # [{"start": float, "end": float, "content": str}, ...]
        segments = await transcribe_audio(str(p))

        lines: list[str] = []

        for seg in segments:
            start, end, content = seg.get("start"), seg.get("end"), seg.get("content", "")

            if content is None:
                continue

            content = str(content)

            if start is not None and end is not None:
                try:
                    lines.append(f"[{float(start):.2f}-{float(end):.2f}] {content}")
                except (TypeError, ValueError):
                    # start/end are broken maybe, just dumping content
                    lines.append(content)
            else:
                lines.append(content)

        return "\n".join(lines).strip()

    # plain text files
    if ext in text_exts:
        return p.read_text(encoding="utf-8", errors="ignore")

    # unsupported
    raise ValueError(f"Unsupported file type for transcription: {ext}")


if __name__ == "__main__":
    import asyncio
    import json
    from pathlib import Path
    from tqdm import tqdm

    CONCURRENCY = 4
    PATH = "/workspace/userspace/telegram-smart-search/smart-search-engine/resources/privacy_dataset.jsonl"
    OUT_PATH = str(Path(PATH).parent / "output.jsonl")

    async def main() -> None:
        async def run_one(text: str, is_mask_needed: bool) -> tuple[bool, str]:
            masked_text = await mask_text(text)
            ok = ("*" in masked_text) == is_mask_needed
            return ok, masked_text

        # total lines for tqdm
        with open(PATH, "r", encoding="utf-8") as f:
            total_lines = sum(1 for _ in f)

        answered_correctly = 0
        total_count = 0

        pbar = tqdm(total=total_lines, desc="Masking", unit="line")

        try:
            async def worker(item: tuple[str, bool]) -> tuple[bool, str, bool]:
                text, is_mask_needed = item
                ok, masked_text = await run_one(text, is_mask_needed)
                return ok, masked_text, is_mask_needed

            with open(PATH, "r", encoding="utf-8") as fin, open(OUT_PATH, "w", encoding="utf-8") as fout:
                in_flight: set[asyncio.Task[tuple[bool, str, bool]]] = set()

                async def drain_finished() -> None:
                    nonlocal answered_correctly, total_count
                    done, _ = await asyncio.wait(in_flight, timeout=0, return_when=asyncio.FIRST_COMPLETED)
                    for task in done:
                        in_flight.remove(task)
                        ok, masked_text, is_mask_needed = task.result()
                        answered_correctly += int(ok)
                        total_count += 1
                        # note: text is captured via closure below in task creation (see payload)
                        payload = task._payload
                        fout.write(
                            json.dumps(
                                {
                                    "text": payload["text"],
                                    "is_mask_needed": payload["is_mask_needed"],
                                    "masked_text": masked_text,
                                },
                                ensure_ascii=False,
                            )
                            + "\n"
                        )

                for line in fin:
                    pbar.update(1)
                    line = line.strip()
                    if not line:
                        continue

                    data = json.loads(line)
                    text = data["text"]
                    is_mask_needed = bool(data["mask"])

                    # backpressure: keep at most CONCURRENCY tasks in flight
                    while len(in_flight) >= CONCURRENCY:
                        await drain_finished()

                    t = asyncio.create_task(worker((text, is_mask_needed)))
                    # attach payload for writing output later (avoid keeping separate dicts/queues)
                    t._payload = {"text": text, "is_mask_needed": is_mask_needed}
                    in_flight.add(t)

                # finish remaining tasks
                while in_flight:
                    await drain_finished()

        finally:
            pbar.close()

        print(round(answered_correctly / total_count, 4))
        print(f"Wrote: {OUT_PATH}")

    asyncio.run(main())